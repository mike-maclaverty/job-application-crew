# Warning control
import warnings
warnings.filterwarnings('ignore')

import os
import requests
from dotenv import load_dotenv
import streamlit as st
from crewai import Agent, Task, Crew
from docx import Document
import zipfile
import io
import tempfile

load_dotenv()

# Function to read job description from URL
def get_job_description(url):
    response = requests.get(url)
    return response.text

# Function to read LinkedIn profile from URL
def get_linkedin_profile(url):
    response = requests.get(url)
    return response.text

# Function to read the resume docx file and return the text and file path
def read_resume(file):
    doc = Document(file)
    resume_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    # Write the Markdown content to a file in a temporary directory
    with tempfile.NamedTemporaryFile(delete=False, suffix='.md', dir=os.getcwd()) as temp_file:
        output_file = temp_file.name
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(resume_text)
    return resume_text, output_file

st.title("Resume Customization App")

# Input fields
job_description_url = st.text_input("Enter Job Description URL:")
linkedin_profile_url = st.text_input("Enter LinkedIn Profile URL:")
openai_api_key = st.text_input("Enter OpenAI API Key:", type="password")
serper_api_key = st.text_input("Enter SERPER API Key:", type="password")
resume_file = st.file_uploader("Upload your resume (.docx)", type=["docx"])
github_profile_url = ""
github_profile_url = st.text_input("Enter GitHub Profile URL (optional):")

resume_text = ""
resume_file_path = ""

if st.button("Customize Resume"):
    if job_description_url and linkedin_profile_url and openai_api_key and serper_api_key and resume_file:
        
        os.environ["OPENAI_API_KEY"] = openai_api_key
        os.environ["OPENAI_MODEL_NAME"] = 'gpt-4o'
        
        # Process job description
        job_description = get_job_description(job_description_url)
        # Process LinkedIn profile
        linkedin_profile = get_linkedin_profile(linkedin_profile_url)
        # Process resume
        resume_text, resume_file_path = read_resume(resume_file)

        #Create tools
        from crewai_tools import FileReadTool, SerperDevTool, ScrapeWebsiteTool
        
        search_tool = SerperDevTool()
        scrape_tool = ScrapeWebsiteTool()
        read_resume_tool = FileReadTool(file_path=resume_file_path)                      
        
        ##Create agents
        # Agent 1: Researcher
        researcher = Agent(
            role="Professional Job Researcher",
            goal="ALWAYS ONLY USE THE JOB POSTING PROVIDED (i.e. job_description & job_description_url)!!!"
                 "DO NOT USE ANY OTHER JOB POSTINGS OR INFORMATION."
                 "DO NOT TRY TO SEARCH FOR ANY OTHER SIMILAR JOBS, JUST USE THE JOB POSTING PROVIDED!!!"
                 "Make sure to do amazing analysis on "
                 "the supplied job posting to help job applicants",
            tools = [scrape_tool],
            verbose=True,
            backstory=(
                "As a Professional Job Researcher, your prowess in "
                "navigating and extracting critical "
                "information from job postings is unmatched."
                "Your skills help pinpoint the necessary "
                "qualifications and skills sought "
                "by employers, forming the foundation for "
                "effective application tailoring."
            )
        )
        # Agent 2: Profiler
        profiler = Agent(
            role="Personal Profiler for job applicants",
            goal="Do increditble research on job applicants "
                "to help them stand out in the job market"
                "Leverage the resume, LinkedIn profile, and GitHub (if available) to create a comprehensive profile.",
            tools = [scrape_tool, read_resume_tool],
            verbose=True,
            backstory=(
                "Equipped with analytical prowess, you dissect "
                "and synthesize information "
                "from diverse sources to craft comprehensive "
                "personal and professional profiles, laying the "
                "groundwork for personalized resume enhancements."
            )
        )
        # Agent 3: Resume Strategist
        resume_strategist = Agent(
            role="Senior Resume Strategist",
            goal="Find all the best ways to make a "
                "resume stand out in the job market."
                "Ensure the resume is tailored to the job requirements."
                "ALWAYS make sure you are only including the formatted resume in the output file, not your notes.",
            tools = [scrape_tool, search_tool,
                    read_resume_tool],
            verbose=True,
            backstory=(
                "With a strategic mind and an eye for detail, you "
                "excel at refining resumes to highlight the most "
                "relevant skills and experiences, ensuring they "
                "resonate perfectly with the job's requirements."
            )
        )
        # Agent 4: Interview Preparer
        interview_preparer = Agent(
            role="Engineering Interview Preparer",
            goal="Create interview questions and talking points "
                "based on the resume and job requirements",
            tools = [scrape_tool, search_tool,
                    read_resume_tool],
            verbose=True,
            backstory=(
                "Your role is crucial in anticipating the dynamics of "
                "interviews. With your ability to formulate key questions "
                "and talking points, you prepare candidates for success, "
                "ensuring they can confidently address all aspects of the "
                "job they are applying for."
            )
        )

        ##Create tasks
        # Task for Researcher Agent: Extract Job Requirements
        research_task = Task(
            description=(
                "ALWAYS ONLY USE THE JOB POSTING PROVIDED (i.e. job_description & job_description_url)!!!"
                "DO NOT USE ANY OTHER JOB POSTINGS OR INFORMATION."
                "DO NOT TRY TO SEARCH FOR ANY OTHER SIMILAR JOBS, JUST USE THE JOB POSTING PROVIDED!!!"
                "Analyze ONLY the job description provided ({job_description}) "
                "to extract key skills, experiences, and qualifications "
                "required. Use the tools to gather content and identify "
                "and categorize the requirements."
            ),
            expected_output=(
                "A structured list of job requirements, including necessary "
                "skills, qualifications, and experiences. ALWAYS ONLY USE THE JOB POSTING PROVIDED (i.e. job_description)!!!"
            ),
            agent=researcher,
            async_execution=True
        )
        # Task for Profiler Agent: Compile Comprehensive Profile
        profile_task = Task(
            description=(
                "Compile a detailed personal and professional profile "
                "using the provided resume, the LinkedIn profile ({linkedin_profile}), "
                "the GitHub repository ({github_url}). Utilize tools to extract and "
                "synthesize information from these sources."
                "Do not worry if you cannot authenticate on LinkedIn, just pull down and use whatever you can."
            ),
            expected_output=(
                "A comprehensive profile document that includes skills, "
                "project experiences, contributions, interests, and "
                "communication style."
            ),
            agent=profiler,
            async_execution=True
        )
        # Task for Resume Strategist Agent: Align Resume with Job Requirements
        resume_strategy_task = Task(
            description=(
                "Using the profile and job description and requirements context obtained from "
                "previous tasks, tailor the resume to highlight the most "
                "relevant areas. Employ tools to adjust and enhance the "
                "resume content. Make sure this is the best resume ever but "
                "don't make up any information. Update every section, "
                "inlcuding the initial summary, work experience, skills, "
                "and education. All to better reflrect the candidates "
                "abilities and how it matches the job posting."
                "ALWAYS make sure you are only including the formatted resume in the output file, not your notes."
                ),
            expected_output=(
                "An updated resume that effectively highlights the candidate's "
                "qualifications and experiences relevant to the job."
                "ALWAYS make sure you are only including the formatted resume in the output file, not your notes."
            ),
            output_file="tailored_resume.md",
            context=[research_task, profile_task],
            agent=resume_strategist
        )
        # Task for Interview Preparer Agent: Develop Interview Materials
        interview_preparation_task = Task(
            description=(
                "Create a set of potential interview questions and talking "
                "points based on the tailored resume and job requirements. "
                "Utilize tools to generate relevant questions and discussion "
                "points. Make sure to use these question and talking points to "
                "help the candiadte highlight the main points of the resume "
                "and how it matches the job posting."
            ),
            expected_output=(
                "A document containing key questions and talking points "
                "that the candidate should be prepared for the initial interview."
                "Make sure you are including all the formatted interview materials in the output file."
            ),
            output_file="interview_materials.md",
            context=[research_task, profile_task, resume_strategy_task],
            agent=interview_preparer
        )

        #Create the crew
        job_application_crew = Crew(
            agents=[researcher,
                    profiler,
                    resume_strategist,
                    interview_preparer],

            tasks=[research_task,
                profile_task,
                resume_strategy_task,
                interview_preparation_task],
             verbose=True
        )

        #Run the crew
        job_application_inputs = {
            'job_description': job_description,
            'linkedin_profile': linkedin_profile,
            'github_url': github_profile_url
            }

        result = job_application_crew.kickoff(inputs=job_application_inputs)

        # Function to convert Markdown to DOCX
        def convert_md_to_docx(md_file, output_docx):
            doc = Document()

            with open(md_file, "r", encoding="utf-8", errors="ignore") as file:
                lines = file.readlines()
                for line in lines:
                    doc.add_paragraph(line.strip())

            doc.save(output_docx)

        # Convert tailored_resume.md to DOCX
        convert_md_to_docx("tailored_resume.md", "tailored_resume.docx")

        # Convert interview_materials.md to DOCX
        convert_md_to_docx("interview_materials.md", "interview_materials.docx")

        # Create a bytes buffer for the zip file
        zip_buffer = io.BytesIO()

        # Create a zip file in the buffer
        with zipfile.ZipFile(zip_buffer, "w") as zip_file:
            zip_file.write("tailored_resume.docx")
            zip_file.write("interview_materials.docx")

        # Rewind the buffer's file pointer
        zip_buffer.seek(0)

        # Provide the download button for the zip file
        st.download_button(
            label="Download All Files",
            data=zip_buffer,
            file_name="customized_resume_and_interview_materials.zip",
            mime="application/zip"
        )

        # Remove temporary DOCX files
        os.remove("tailored_resume.docx")
        os.remove("interview_materials.docx")
        os.remove("tailored_resume.md")
        os.remove("interview_materials.md")        
        os.remove(resume_file_path)
    else:
        st.write("Missing required inputs: job description URL, LinkedIn profile URL, resume file, OpenAI API key, or SERPER API key.")